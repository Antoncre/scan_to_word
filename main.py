import os
from PIL import Image, ImageFilter, ImageEnhance
import pytesseract
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Jeśli używasz Windows, ustaw ścieżkę do Tesseract
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Ścieżki
input_folder = "images/"
output_file = "OCR_wynik_POL.docx"

# Utwórz dokument Word
doc = Document()

# Lista plików graficznych
supported_ext = ('.jpg', '.jpeg', '.png', '.bmp', '.tif', '.tiff')

for file in sorted(os.listdir(input_folder)):
    if file.lower().endswith(supported_ext):
        path = os.path.join(input_folder, file)
        print(f"📄 Przetwarzam: {file}")

        try:
            image = Image.open(path)

            # 🧠 Poprawa jakości przed OCR
            image = image.convert('L')  # skala szarości
            image = image.filter(ImageFilter.MedianFilter())
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(2.0)  # zwiększ kontrast

            # OCR w języku polskim
            text = pytesseract.image_to_string(image, lang='pol')

            # Dodaj tekst do dokumentu
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text.strip())
            run.font.size = Pt(12)

            doc.add_page_break()

        except Exception as e:
            print(f"❌ Błąd w pliku {file}: {e}")

# Zapisz dokument
doc.save(output_file)
print(f"✅ Dokument zapisany jako: {output_file}")
